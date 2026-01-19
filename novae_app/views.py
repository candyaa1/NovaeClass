from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import LoginView
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.utils import timezone
from django import forms
from django.forms import formset_factory
from django .db import transaction

from django.views.decorators.cache import never_cache
from datetime import timedelta, date
import random

from docx import Document

from .models import (
    User,
    StudentProfile,
    ParentProfile,
    Assignment,
    AssignmentInstance,
    Material,
    Question,
    Game,
    StudyPlan,
    StudentAnswer,
)

from .forms import StudyPlanForm, AssignmentSubmissionForm


# ---------------------------
# ACCESS HELPERS
# ---------------------------
def user_is_paid(user):
    return (
        user.is_authenticated
        and hasattr(user, "billing_profile")
        and user.billing_profile.is_paid
    )


# ---------------------------
# LANDING PAGE
# ---------------------------
def landing(request):
    return render(request, 'novae_app/landing.html')


# ---------------------------
# STUDENT LOGIN
# ---------------------------
class StudentLoginView(LoginView):
    template_name = 'novae_app/student_login.html'

    def form_valid(self, form):
        user = form.get_user()
        if not user.is_student():
            return redirect('landing')
        return super().form_valid(form)


# ---------------------------
# STUDENT DASHBOARD
# ---------------------------
@login_required
def student_dashboard(request):
    if not hasattr(request.user, 'student_profile'):
        return redirect('landing')

    student = request.user.student_profile
    paid = user_is_paid(request.user)

    now = timezone.now()
    today = now.date()

    session_start = request.session.get('session_start', now.timestamp())
    elapsed_seconds = now.timestamp() - session_start

    if student.last_active_date != today:
        student.daily_time_seconds = 0
        student.last_active_date = today

    student.daily_time_seconds += int(elapsed_seconds)
    student.last_active_date = today
    student.save()
    request.session['session_start'] = now.timestamp()

    # ------------------------
    # Choose assignments based on demo vs paid
    # ------------------------
    if paid:
        assignments_queryset = Assignment.objects.all()
    else:
        assignments_queryset = Assignment.objects.filter(is_demo=True)

    # Ensure AssignmentInstances exist
    for assignment in assignments_queryset:
        AssignmentInstance.objects.get_or_create(
            student=student,
            assignment=assignment
        )

    # ------------------------
    # Fetch assignments / grades / materials
    # ------------------------
    if paid:
        assignments = AssignmentInstance.objects.filter(
            student=student, completed=False
        ).order_by('assignment__due_date')

        grades = AssignmentInstance.objects.filter(
            student=student, completed=True, score__isnull=False
        )

        materials = Material.objects.filter(grade_level=student.grade)
    else:
        assignments = AssignmentInstance.objects.filter(
            student=student,
            assignment__is_demo=True,
            completed=False
        ).order_by('assignment__due_date')

        grades = AssignmentInstance.objects.filter(
            student=student,
            assignment__is_demo=True,
            completed=True,
            score__isnull=False
        )

        materials = Material.objects.filter(
            grade_level=student.grade,
            is_demo=True
        )

    return render(request, 'novae_app/student_dashboard.html', {
        'assignments': assignments,
        'grades': grades,
        'materials': materials,
        'demo': not paid,
        'total_time_spent': timedelta(seconds=student.daily_time_seconds),
    })


# ---------------------------
# STUDENT ASSIGNMENTS
# ---------------------------
@login_required
def student_assignments(request):
    if not user_is_paid(request.user):
        messages.warning(request, "Upgrade to access assignments.")
        return redirect('billing')

    student = request.user.student_profile
    instances = AssignmentInstance.objects.filter(student=student)
    return render(request, 'novae_app/student_assignments.html', {
        'assignments': instances
    })

@login_required
def student_assignment_retake(request, instance_id):
    instance = get_object_or_404(AssignmentInstance, id=instance_id, student=request.user.student_profile)
    if not instance.retake_allowed():
        return redirect('student_assignments')
    instance.start_retake()
    return redirect('student_assignment_detail', instance_id=instance.id)

# ---------------------------
# ASSIGNMENT DETAIL
# ---------------------------
@login_required
def student_assignment_detail(request, instance_id):
    if not user_is_paid(request.user):
        return redirect('billing')

    student = request.user.student_profile
    instance = get_object_or_404(
        AssignmentInstance,
        id=instance_id,
        student=student
    )

    questions = instance.assignment.questions.all()

    if request.method == 'POST':
        correct = 0
        total = questions.count()

        for question in questions:
            value = request.POST.get(f'question_{question.id}', '').strip()

            answer, _ = StudentAnswer.objects.update_or_create(
                student=student,
                question=question,
                assignment_instance=instance
            )

            if question.question_type == 'TEXT':
                answer.text_answer = value
                if value.lower() == (question.correct_option or '').lower():
                    correct += 1
            else:
                answer.selected_option = value
                if value == question.correct_option:
                    correct += 1

            answer.save()

        instance.score = (correct / total) * 100 if total else 0
        instance.completed = True
        instance.save()

        return redirect('student_assignments')

    return render(request, 'novae_app/assignment_detail.html', {
        'instance': instance,
        'questions': questions,
    })


# ---------------------------
# STUDENT MATERIALS
# ---------------------------
@login_required
def student_materials(request):
    if not user_is_paid(request.user):
        return redirect('billing')

    student = request.user.student_profile
    materials = Material.objects.filter(grade_level=student.grade)
    return render(request, 'novae_app/student_materials.html', {'materials': materials})


# ---------------------------
# DAILY QUIZ
# ---------------------------
@login_required
def daily_quiz(request):
    questions = list(Question.objects.all())

    if not user_is_paid(request.user):
        questions = questions[:1]

    if not questions:
        return render(request, 'novae_app/daily_quiz.html', {'error': "No questions available."})

    question = random.choice(questions)

    if request.method == 'POST':
        selected_option = request.POST.get('option')
        correct = (selected_option == question.correct_option)
        return render(request, 'novae_app/daily_quiz_result.html', {
            'question': question,
            'selected_option': selected_option,
            'correct': correct,
        })

    return render(request, 'novae_app/daily_quiz.html', {'question': question})


# ---------------------------
# LEARNING GAMES
# ---------------------------
@login_required
def learning_games_view(request):
    if not user_is_paid(request.user):
        return redirect('billing')

    grade_map = {
        'K': 0, '1st': 1, '2nd': 2, '3rd': 3, '4th': 4,
        '5th': 5, '6th': 6, '7th': 7, '8th': 8,
        '9th': 9, '10th': 10, '11th': 11, '12th': 12
    }

    grade = grade_map.get(request.user.student_profile.grade, 0)
    games = Game.objects.filter(min_grade__lte=grade, max_grade__gte=grade)
    return render(request, 'novae_app/learning_games.html', {'games': games})


# ---------------------------
# STUDY PLANS
# ---------------------------
@login_required
def study_plan_list(request):
    if not user_is_paid(request.user):
        return redirect('billing')

    plans = StudyPlan.objects.filter(user=request.user)
    return render(request, 'novae_app/study_plan_list.html', {'plans': plans})


@login_required
def study_plan_create(request):
    if not user_is_paid(request.user):
        return redirect('billing')

    if request.method == 'POST':
        form = StudyPlanForm(request.POST)
        if form.is_valid():
            plan = form.save(commit=False)
            plan.user = request.user
            plan.save()
            return redirect('study_plan_list')
    else:
        form = StudyPlanForm()

    return render(request, 'novae_app/study_plan_form.html', {'form': form})


# ---------------------------
# DOWNLOAD ASSIGNMENTS
# ---------------------------
@login_required
def download_assignment_docx(request, pk):
    if not user_is_paid(request.user):
        return redirect('billing')

    assignment = get_object_or_404(Assignment, id=pk)

    doc = Document()
    doc.add_heading(assignment.title, level=1)
    doc.add_paragraph(f"Due Date: {assignment.due_date}")
    doc.add_paragraph(assignment.description or "No description provided.")

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{assignment.title}.docx"'
    doc.save(response)
    return response
@login_required
def student_grades(request):
    student = request.user.student_profile
    grades = AssignmentInstance.objects.filter(student=student, score__isnull=False)
    return render(request, 'novae_app/student_grades.html', {'grades': grades})

@login_required
def get_grades(request, child_id):
    try:
        student = StudentProfile.objects.get(user__id=child_id)
    except StudentProfile.DoesNotExist:
        return JsonResponse({"error": "Student not found"}, status=404)
    
    grades = AssignmentInstance.objects.filter(student=student, score__isnull=False)
    
    grades_data = [{
        'assignment': g.assignment.title,
        'score': g.score,
        'comments': g.feedback or 'No feedback available',
    } for g in grades]
    
    return JsonResponse({'grades': grades_data})


@login_required
def parent_submitted_assignments(request, student_id):
    parent = request.user.parent_profile
    student = get_object_or_404(StudentProfile, id=student_id, parents=parent)
    assignments = AssignmentInstance.objects.filter(student=student, score__isnull=False)
    return render(request, 'novae_app/student_grades.html', {'student': student, 'grades': assignments})
@never_cache
def student_signup(request):
    if request.method == 'POST':
        parent_form = ParentSignUpForm(request.POST)
        child_formset = ChildFormSet(request.POST)
        if parent_form.is_valid() and child_formset.is_valid():
            with transaction.atomic():
                parent_user = User.objects.create_user(
                    username=parent_form.cleaned_data['username'],
                    email=parent_form.cleaned_data['email'],
                    password=parent_form.cleaned_data['password'],
                    role='parent'
                )
                parent_profile = ParentProfile.objects.create(user=parent_user)
                for child_form in child_formset:
                    username = child_form.cleaned_data.get('username')
                    grade = child_form.cleaned_data.get('grade')
                    password = child_form.cleaned_data.get('password')
                    if username and grade and password:
                        child_user = User.objects.create_user(username=username, password=password, role='student')
                        student_profile = StudentProfile.objects.create(user=child_user, grade=grade)
                        parent_profile.children.add(student_profile)
            return redirect('landing')
    else:
        parent_form = ParentSignUpForm()
        child_formset = ChildFormSet()

    return render(request, "student_signup.html", {
      'parent_form': parent_form,
      'child_formset': child_formset,
})

    
def submit_assignment(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    student = get_object_or_404(StudentProfile, user=request.user)
    questions = assignment.questions.all()
    if request.method == "POST":
        form = AssignmentSubmissionForm(questions, request.POST)
        if form.is_valid():
            for question in questions:
                answer_value = form.cleaned_data.get(f'q_{question.id}')
                StudentAnswer.objects.create(
                    student=student,
                    question=question,
                    answer_text=answer_value if question.is_text_answer else None,
                    selected_option=answer_value if not question.is_text_answer else None
                )
            return redirect('assignment_success')
def coming_soon(request):
    return render(request, "coming_soon.html")

def about_us(request):
    return render(request, 'about_us.html')


# ---------------------------
# PARENT DASHBOARD
# ---------------------------
@login_required
def parent_dashboard(request):
    if not user_is_paid(request.user):
        return redirect('billing')

    if not request.user.is_parent():
        return redirect('landing')

    children = request.user.parent_profile.children.all()
    data = []

    today = timezone.now().date()

    for child in children:
        grades = AssignmentInstance.objects.filter(
            student=child,
            score__isnull=False
        )

        seconds = child.daily_time_seconds if child.last_active_date == today else 0

        data.append({
            'student_id': child.id,
            'user': child.user,
            'grade': child.grade,
            'grades': grades,
            'total_time_spent': timedelta(seconds=seconds),
        })

    return render(request, 'novae_app/parent_dashboard.html', {'children': data})
@login_required
def study_plan_edit(request, pk):
    if not user_is_paid(request.user):
        return redirect('billing')

    plan = get_object_or_404(StudyPlan, id=pk, user=request.user)

    if request.method == 'POST':
        form = StudyPlanForm(request.POST, instance=plan)
        if form.is_valid():
            form.save()
            return redirect('study_plan_list')
    else:
        form = StudyPlanForm(instance=plan)

    return render(request, 'novae_app/study_plan_form.html', {'form': form})


@login_required
def study_plan_delete(request, pk):
    if not user_is_paid(request.user):
        return redirect('billing')

    plan = get_object_or_404(StudyPlan, id=pk, user=request.user)

    if request.method == 'POST':
        plan.delete()
        return redirect('study_plan_list')

    return render(request, 'novae_app/study_plan_confirm_delete.html', {'plan': plan})



# ---------------------------
# DOWNLOAD GRADED ASSIGNMENT DOCX
# ---------------------------
@login_required
def download_graded_assignment_docx(request, instance_id):
    instance = get_object_or_404(AssignmentInstance, id=instance_id)

    doc = Document()
    doc.add_heading(f"{instance.assignment.title} - Graded Review", level=1)
    doc.add_paragraph(f"Due Date: {instance.assignment.due_date}")
    doc.add_paragraph(f"Score: {instance.score or 'Not graded yet'}")
    doc.add_paragraph("Description:")
    doc.add_paragraph(instance.assignment.description or "No description provided.")
    doc.add_paragraph("Feedback:")
    doc.add_paragraph(instance.feedback or "No feedback provided.")

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"{instance.assignment.title}_graded_review.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response


# ---------------------------
# PARENT LOGIN VIEW
# ---------------------------
class ParentLoginView(LoginView):
    template_name = 'novae_app/parent_login.html'

    def form_valid(self, form):
        user = form.get_user()
        # Make sure the user is a parent
        if user.is_authenticated and hasattr(user, 'is_parent') and user.is_parent():
            return super().form_valid(form)
        form.add_error(None, "You must log in as a parent.")
        return self.form_invalid(form)

    def get_success_url(self):
        return '/parent/dashboard/'

@login_required
def achievements(request):
    """Render the achievements page for the logged-in user."""
    return render(request, 'novae_app/achievements.html')

@login_required
def study_timer(request):
    """Render the study timer page for the logged-in user."""
    return render(request, 'novae_app/study_timer.html')

@login_required
def assignment_results(request, child_name):
    """
    Render assignment results page for a specific student.
    """
    # Fetch the student user
    student_user = get_object_or_404(User, username=child_name)

    # Fetch student profile
    student = get_object_or_404(StudentProfile, user=student_user)

    # Fetch all AssignmentInstances for this student
    assignments_instances = AssignmentInstance.objects.filter(student=student)

    if not assignments_instances.exists():
        return render(request, 'novae_app/assignment_results.html', {'error': 'No assignments found'})

    assignments_data = []

    for instance in assignments_instances:
        questions = instance.assignment.questions.all()  # Assuming related_name='questions' in Assignment
        answers = Answer.objects.filter(assignment_instance=instance)  # Assuming Answer model exists

        assignment_data = {
            'assignment': instance.assignment,
            'score': instance.score,
            'submitted_on': instance.submitted_on,
            'questions': []
        }

        for question in questions:
            student_answer = answers.filter(question=question).first()
            assignment_data['questions'].append({
                'question_text': question.text,
                'correct_answer': question.correct_option,
                'student_answer': student_answer.answer_text if student_answer else 'Not answered',
                'is_correct': student_answer.is_correct if student_answer else False,
            })

        assignments_data.append(assignment_data)

    return render(request, 'novae_app/assignment_results.html', {'assignments_data': assignments_data})

class ParentSignUpForm(forms.Form):
    """Form for parent signup."""
    username = forms.CharField(
        max_length=150,
        label="Username",
        widget=forms.TextInput(attrs={'placeholder': 'Username'})
    )
    email = forms.EmailField(
        label="Email",
        widget=forms.EmailInput(attrs={'placeholder': 'Email'})
    )
    password = forms.CharField(
        label="Password",
        widget=forms.PasswordInput(attrs={'placeholder': 'Password'})
    )

class ChildForm(forms.Form):
    """Form for adding a child account."""
    username = forms.CharField(
        max_length=150,
        label="Child Username",
        widget=forms.TextInput(attrs={'placeholder': 'Child Username'})
    )
    grade_choices = [
        ('K', 'K'),
        ('1st', '1st'),
        ('2nd', '2nd'),
        ('3rd', '3rd'),
        ('4th', '4th'),
        ('5th', '5th'),
        ('6th', '6th'),
        ('7th', '7th'),
        ('8th', '8th'),
        ('9th', '9th'),
        ('10th', '10th'),
        ('11th', '11th'),
        ('12th', '12th'),
    ]
    grade = forms.ChoiceField(
        choices=grade_choices,
        label="Grade"
    )
    password = forms.CharField(
        label="Password",
        widget=forms.PasswordInput(attrs={'placeholder': 'Password'})
    )


# Formset to allow adding multiple children
ChildFormSet = formset_factory(ChildForm, extra=1)

@login_required
def student_assignment_retake(request, instance_id):
    """
    Allow a student to retake an assignment if retake is allowed.
    """
    student = request.user.student_profile
    instance = get_object_or_404(AssignmentInstance, id=instance_id, student=student)

    if not instance.retake_allowed():
        return redirect('student_assignments')

    instance.start_retake()
    return redirect('student_assignment_detail', instance_id=instance.id)


# novae_app/views.py
from django.shortcuts import render
from .models import BillingProfile, Course, Material

def billing_view(request):
    user = request.user
    billing_profile = None
    free_trial_items = []

    if hasattr(user, 'billingprofile'):
        billing_profile = user.billingprofile
        # If the user is on a free trial, show free trial items
        if billing_profile.plan == 'free_trial':
            # Example: show free courses or free materials
            free_trial_items = Course.objects.filter(is_demo=True) | Material.objects.filter(is_demo=True)

    context = {
        'billing_profile': billing_profile,
        'free_trial_items': free_trial_items,
    }
    return render(request, 'billing.html', context)
